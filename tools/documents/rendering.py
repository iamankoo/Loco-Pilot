"""The actual reportlab/python-docx/openpyxl/pypdf calls behind
`tools/documents/tools.py` — kept separate from the `Tool` classes so the
rendering logic is directly unit-testable without going through the full
Tool/ToolContext/workspace machinery.
"""

from __future__ import annotations

import csv
import re
from pathlib import Path

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from tools.documents.schemas import CellValue, DocumentSection, SpreadsheetSheet

_STYLES = getSampleStyleSheet()


def _escape(text: str) -> str:
    # reportlab's Paragraph interprets a small subset of HTML-like markup in
    # its text — escape the characters that would otherwise be misread as
    # markup, so arbitrary generated content always renders as plain text.
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _paragraphs_for(text: str) -> list[Paragraph]:
    blocks = []
    for para in text.split("\n\n"):
        if para.strip():
            blocks.append(Paragraph(_escape(para).replace("\n", "<br/>"), _STYLES["BodyText"]))
            blocks.append(Spacer(1, 0.12 * inch))
    return blocks


def render_pdf(path: Path, *, title: str | None, sections: list[DocumentSection], body: str) -> int:
    story: list = []
    if title:
        story.append(Paragraph(_escape(title), _STYLES["Title"]))
        story.append(Spacer(1, 0.25 * inch))

    if sections:
        for section in sections:
            if section.heading:
                story.append(Paragraph(_escape(section.heading), _STYLES["Heading2"]))
                story.append(Spacer(1, 0.1 * inch))
            story.extend(_paragraphs_for(section.body))
    elif body:
        story.extend(_paragraphs_for(body))

    if not story:
        story.append(Paragraph("", _STYLES["BodyText"]))

    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    doc.build(story)
    return doc.page  # updated in place during build(); holds the final page count after it returns


def render_image_as_pdf(path: Path, image_path: Path) -> int:
    from PIL import Image as PILImage

    with PILImage.open(image_path) as img:
        width, height = img.size
    max_w, max_h = LETTER[0] - inch, LETTER[1] - inch
    scale = min(max_w / width, max_h / height, 1.0)

    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    story = [RLImage(str(image_path), width=width * scale, height=height * scale)]
    doc.build(story)
    return doc.page


# Best-effort, line-based Markdown -> reportlab blocks — headings, bullet/
# numbered list lines, and inline **bold**/*italic* — not a claim of
# implementing the full CommonMark spec (tables, nested lists, code fences,
# and links are rendered as plain text rather than parsed).
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_ITALIC_RE = re.compile(r"(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)")


def _inline_markdown_to_reportlab(text: str) -> str:
    escaped = _escape(text)
    escaped = _BOLD_RE.sub(r"<b>\1</b>", escaped)
    escaped = _ITALIC_RE.sub(r"<i>\1</i>", escaped)
    return escaped


def render_markdown_as_pdf(path: Path, markdown_text: str) -> int:
    story: list = []
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            story.append(Spacer(1, 0.08 * inch))
            continue
        heading = _HEADING_RE.match(stripped)
        bullet = _BULLET_RE.match(stripped)
        if heading:
            level = len(heading.group(1))
            style = "Heading1" if level <= 2 else "Heading2" if level <= 4 else "Heading3"
            story.append(Paragraph(_inline_markdown_to_reportlab(heading.group(2)), _STYLES[style]))
        elif bullet:
            story.append(Paragraph("&bull; " + _inline_markdown_to_reportlab(bullet.group(1)), _STYLES["BodyText"]))
        else:
            story.append(Paragraph(_inline_markdown_to_reportlab(stripped), _STYLES["BodyText"]))

    if not story:
        story.append(Paragraph("", _STYLES["BodyText"]))

    doc = SimpleDocTemplate(str(path), pagesize=LETTER)
    doc.build(story)
    return doc.page


def render_docx(path: Path, *, title: str | None, sections: list[DocumentSection], body: str) -> None:
    from docx import Document

    doc = Document()
    if title:
        doc.add_heading(title, level=0)

    if sections:
        for section in sections:
            if section.heading:
                doc.add_heading(section.heading, level=1)
            for para in section.body.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para)
    elif body:
        for para in body.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para)

    doc.save(str(path))


def render_xlsx(path: Path, sheets: list[SpreadsheetSheet]) -> tuple[int, int]:
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    total_rows = 0
    for sheet_spec in sheets:
        title = (sheet_spec.name or "Sheet1")[:31]  # Excel's own 31-char sheet-name limit
        ws = wb.create_sheet(title=title)
        if sheet_spec.headers:
            ws.append(sheet_spec.headers)
        for row in sheet_spec.rows:
            ws.append(row)
        total_rows += len(sheet_spec.rows)
    wb.save(str(path))
    return len(sheets), total_rows


def render_csv(path: Path, *, headers: list[str], rows: list[list[CellValue]]) -> None:
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.writer(fh)
        if headers:
            writer.writerow(headers)
        writer.writerows(rows)


def extract_pdf_text(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def read_xlsx_first_sheet(path: Path) -> list[list[object]]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    try:
        ws = wb[wb.sheetnames[0]]
        return [list(row) for row in ws.iter_rows(values_only=True)]
    finally:
        wb.close()
